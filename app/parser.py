import fitz  # PyMuPDF
from docx import Document as DocxDocument
import os
import logging, pytesseract
from pdfminer.high_level import extract_text as pdfminer_extract
from pdf2image import convert_from_path


# Configure logging for better debugging
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')

def extract_text_from_pdf(filepath):
    """Extracts text from a PDF file using PyMuPDF (fitz)."""
    text = ""
    try:
        with fitz.open(filepath) as doc:
            for page in doc:
                text += page.get_text()
    except Exception as e:
        logging.error(f"Error opening or reading PDF file {filepath}: {e}")
        return ""
    return text.strip()

from docx import Document as DocxDocument

def extract_text_from_docx(filepath):
    """
    Extracts readable text from a .docx Word file, including paragraphs and table cells.
    """
    try:
        doc = DocxDocument(filepath)
        full_text = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))

        return "\n".join(full_text)
    except Exception as e:
        logging.exception(f"❌ Failed to extract DOCX text: {e}")
        return ""

def extract_text_from_file(filepath):
    """
    Extracts text from PDF, DOCX, or TXT files.
    Falls back to OCR for scanned PDFs.
    """
    if not os.path.exists(filepath):
        logging.warning(f"File not found: {filepath}")
        return ""

    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        try:
            text = pdfminer_extract(filepath)
            if text and len(text.strip()) > 100:
                logging.info(f"PDF text extracted using pdfminer: {len(text)} characters")
                return text
            else:
                logging.warning("PDF text empty or too short — falling back to OCR")
                images = convert_from_path(filepath)
                ocr_text = ""
                for i, img in enumerate(images):
                    ocr_page = pytesseract.image_to_string(img)
                    logging.info(f"OCR page {i+1}: {len(ocr_page)} characters")
                    ocr_text += ocr_page + "\n"
                return ocr_text.strip()
        except Exception as e:
            logging.exception(f"PDF extraction failed: {e}")
            return ""

    elif ext == ".docx":
        return extract_text_from_docx(filepath)

    elif ext == ".txt":
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read().strip()
        except Exception as e:
            logging.error(f"Error reading TXT file {filepath}: {e}")
            return ""

    else:
        logging.warning(f"Unsupported file type: {ext}")
        return ""
 

